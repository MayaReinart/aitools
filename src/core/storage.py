"""Job data storage utilities."""

import json
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any

from docx import Document
from loguru import logger

from src.core.config import settings

# Default path for job data, can be overridden in tests
JOB_DATA_ROOT = settings.job_data_path


class JobArtifact(str, Enum):
    """Types of artifacts that can be stored for a job."""

    SPEC = "spec"
    PARSED_SPEC = "parsed_spec"
    SUMMARY = "summary"
    EXPORT = "export"
    LOG = "log"
    OUTPUT = "output"


class SpecFormat(str, Enum):
    """Formats for API specifications."""

    YAML = "yaml"
    JSON = "json"


class ExportFormat(str, Enum):
    """Formats for exported summaries."""

    MARKDOWN = "md"
    HTML = "html"
    DOCX = "docx"


def _get_and_log_path(path: Path, job_id: str, artifact: str) -> Path | None:
    """Check if a path exists and log the result."""
    if path.exists():
        logger.debug(f"Found {job_id} {artifact} at {path}")
        return path
    logger.debug(f"No {job_id} {artifact} found at {path}")
    return None


class JobStorage:
    """Handles storage and retrieval of job-related data."""

    def __init__(self, job_id: str) -> None:
        """Initialize storage for a job.

        Args:
            job_id: Unique identifier for the job
        """
        self.job_id = job_id
        self.job_dir = JOB_DATA_ROOT / job_id

        # Ensure root directory exists with proper permissions
        JOB_DATA_ROOT.mkdir(mode=0o755, parents=True, exist_ok=True)
        # Create job directory with proper permissions
        self.job_dir.mkdir(mode=0o755, parents=True, exist_ok=True)

        # Initialize log file
        self.log_file = self.job_dir / "log.txt"
        if not self.log_file.exists():
            self.log_file.touch(mode=0o644)

    def save_spec(self, content: str, format_: SpecFormat) -> Path:
        """Save the uploaded spec file.

        Args:
            content: The spec content to save
            format_: The format of the spec (YAML or JSON)

        Returns:
            Path to the saved spec file
        """
        spec_path = self.job_dir / f"spec.{format_.value}"
        spec_path.write_text(content)
        self.log_event("Saved spec file")
        logger.info(f"Saved {self.job_id} spec to {spec_path}")
        return spec_path

    def save_summary(self, summary: dict[str, Any]) -> Path:
        """Save the generated summary."""
        summary_path = self.job_dir / "summary.json"
        summary_path.write_text(json.dumps(summary, indent=2))
        self.log_event("Saved summary")
        logger.info(f"Saved {self.job_id} summary to {summary_path}")
        return summary_path

    def save_export(self, content: str | bytes, format_: ExportFormat) -> Path:
        """Save an exported summary file."""
        export_path = self.job_dir / f"summary.{format_}"
        if isinstance(content, str):
            export_path.write_text(content)
        else:
            export_path.write_bytes(content)
        self.log_event(f"Saved {format_} export")
        logger.info(f"Saved {self.job_id} export to {export_path}")
        return export_path

    def save_parsed_spec(self, parsed_spec: dict[str, Any]) -> Path:
        """Save the parsed OpenAPI spec."""
        parsed_spec_path = self.job_dir / "parsed_spec.json"
        parsed_spec_path.write_text(json.dumps(parsed_spec, indent=2))
        self.log_event("Saved parsed spec")
        logger.info(f"Saved {self.job_id} parsed spec to {parsed_spec_path}")
        return parsed_spec_path

    def log_event(self, message: str) -> None:
        """Log an event to the execution log file.

        Args:
            message: Event message to log
        """
        timestamp = datetime.now(tz=timezone.utc).isoformat()
        log_entry = f"[{timestamp}] {message}\n"
        with self.log_file.open("a") as f:
            f.write(log_entry)

    def get_spec_path(self) -> Path | None:
        """Get the path to the spec file if it exists."""
        for ext in ["yaml", "json"]:
            path = self.job_dir / f"spec.{ext}"
            if ext_path := _get_and_log_path(path, self.job_id, JobArtifact.SPEC):
                return ext_path
        return None

    def get_summary_path(self) -> Path | None:
        """Get the path to the summary file if it exists."""
        path = self.job_dir / "summary.json"
        return _get_and_log_path(path, self.job_id, JobArtifact.SUMMARY)

    def get_export_path(self, format_: ExportFormat) -> Path | None:
        """Get the path to an export file if it exists."""
        path = self.job_dir / f"summary.{format_}"
        return _get_and_log_path(path, self.job_id, JobArtifact.EXPORT)

    def get_parsed_spec_path(self) -> Path | None:
        """Get the path to the parsed spec file if it exists."""
        path = self.job_dir / "parsed_spec.json"
        return _get_and_log_path(path, self.job_id, JobArtifact.PARSED_SPEC)

    def get_log_path(self) -> Path | None:
        """Get the path to the execution log if it exists."""
        return _get_and_log_path(self.log_file, self.job_id, JobArtifact.LOG)

    def ensure_export_exists(self, format_: ExportFormat) -> Path:
        """Ensure an export file exists, creating a placeholder if needed."""
        path = self.job_dir / f"summary.{format_}"
        logger.info(f"Checking if {self.job_id} {format_} export exists at {path}")

        if path.exists():
            return path

        logger.info(f"Creating {self.job_id} {format_} export at {path}")

        if format_ == ExportFormat.MARKDOWN:
            path.write_text("# API Summary\n\nTo be implemented")
        elif format_ == ExportFormat.HTML:
            path.write_text("<h1>API Summary</h1>\n<p>To be implemented</p>")
        elif format_ == ExportFormat.DOCX:
            # Create a minimal DOCX file with a title
            doc = Document()
            doc.add_heading("API Summary", 0)
            doc.add_paragraph("To be implemented")
            doc.save(str(path))

        return path

    def get_export_content(self, format_: ExportFormat) -> tuple[str | bytes, str]:
        """Get export content and media type for a given format.

        Returns:
            A tuple of (content, media_type)
        """
        path = self.ensure_export_exists(format_)

        if format_ == ExportFormat.HTML:
            return path.read_text(), "text/html"
        if format_ == ExportFormat.DOCX:
            return (
                path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        # MARKDOWN
        return path.read_text(), "text/markdown"

    def save_output(self, filename: str, content: str) -> Path:
        """Save an output file.

        Args:
            filename: Name of the output file
            content: Content to save

        Returns:
            Path to the saved file
        """
        output_path = self.job_dir / filename
        output_path.write_text(content)
        self.log_event(f"Saved output file: {filename}")
        logger.info(f"Saved {self.job_id} output to {output_path}")
        return output_path
